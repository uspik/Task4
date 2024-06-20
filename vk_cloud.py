import time

import requests
import docx
from pdf2image import convert_from_bytes
import io
import asyncio
lock = asyncio.Lock()
access_token = "XfC453xRiPDZo1ZXvWU3H5dJM9TdcUEeW1AsSCCP1DSHK9yVB"
token_expired = 0
def get_access_token(client_id, refresh_token):
    url = "https://mcs.mail.ru/auth/oauth/v1/token"
    headers = {
        "Content-Type": "application/json"
    }
    data = {
        "client_id": client_id,
        "refresh_token": refresh_token,
        "grant_type": "refresh_token"
    }
    response = requests.post(url, headers=headers, json=data)
    return response.json()

client_id = "mcs9064514905.ml.vision.3Yzx3KihTHQ5NjsZuVr2w"
client_secret = "B1ckyuc8AXL7QhBNWGqKWz5F4gEFW5wE6kZ4SwnjFSRxbpp67ZbCPh7zYUjD1"
refresh_token = "2JZfPKgXiBL1zqijvk85yS5aadLgvYLUrstURL2EWjDiSMoWFS"
#print(get_access_token(client_id, refresh_token))
async def send_image(file):
    global token_expired, access_token
    if round(time.time() * 1000) >= token_expired:
            async with lock:
                refresh_tok()
    url = f'https://smarty.mail.ru/api/v1/text/recognize?oauth_token={access_token}&oauth_provider=mcs'

    headers = {
        'accept': 'application/json',
    }

    files = {
        'file': (f'MyImage.jpeg', file, 'image/jpeg'),
    }

    data = {
        'meta': '{"images":[{"name":"file"}]}'
    }

    response = requests.post(url, headers=headers, files=files, data=data)
    if response.status_code == 200:
        return response.json()['body']['objects'][0]['text']
    else:
        return "ERROR"

async def getText(file):
    doc = docx.Document(file)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText), len(doc.paragraphs)

async def pdf_to_img(file):
    global access_token
    images = convert_from_bytes(pdf_file=file.read())
    output_str = ''
    for image in images:
        img_byte_arr = io.BytesIO()
        image.save(img_byte_arr, format='PNG')
        img_byte_arr = img_byte_arr.getvalue()
        output_str += await send_image(img_byte_arr, access_token)
    return output_str, len(images)


def refresh_tok():
    global access_token, token_expired
    access_token = get_access_token(client_id, refresh_token)['access_token']
    token_expired = time.time() + 3600
    return

#refresh_tok()

# {'refresh_token': '2JZfPKgXiBL1zqijvk85yS5aadLgvYLUrstURL2EWjDiSMoWFS', 'access_token': 'b5oqfPZC1VSejggxuQtQfukfzhtBy6Nn8U4kTfngqRYG9xDak', 'expired_in': '3600', 'scope': {'objects': 1, 'video': 1, 'persons': 1}}